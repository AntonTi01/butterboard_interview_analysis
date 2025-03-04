import os
import re
from docx import Document

def clean_text(text: str) -> str:
    """Удаляет пустые строки между абзацами."""
    paragraphs = text.strip().split("\n")
    cleaned_paragraphs = [p.strip() for p in paragraphs if p.strip()]
    return "\n".join(cleaned_paragraphs)

def process_srt(input_path: str, output_path: str):
    """Обрабатывает .srt файл, удаляя пустые строки между субтитрами."""
    with open(input_path, "r", encoding="utf-8") as file:
        content = file.read()
    
    cleaned_content = clean_text(content)
    
    with open(output_path, "w", encoding="utf-8") as file:
        file.write(cleaned_content)
    print(f"Обработанный .srt файл сохранён в {output_path}")

def process_docx(input_path: str, output_path: str):
    """Обрабатывает .docx файл, удаляя пустые строки между абзацами."""
    doc = Document(input_path)
    cleaned_paragraphs = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            cleaned_paragraphs.append(text)
    
    new_doc = Document()
    for para in cleaned_paragraphs:
        new_doc.add_paragraph(para)
    
    new_doc.save(output_path)
    print(f"Обработанный .docx файл сохранён в {output_path}")

def process_transcription_file(input_path: str, output_path: str):
    """Обрабатывает отдельный файл .srt или .docx."""
    if input_path.endswith(".srt"):
        process_srt(input_path, output_path)
    elif input_path.endswith(".docx"):
        process_docx(input_path, output_path)
    else:
        print(f"Файл {input_path} не поддерживается (не .srt или .docx)")

def process_speaker_segments(input_file, output_file):
    merged_segments = []

    with open(input_file, "r") as file:
        lines = file.readlines()

    for line in lines:
        parts = line.strip().split()
        start_time = float(parts[0].split("=")[1][:-1])  # Извлекаем число без "s"
        stop_time = float(parts[1].split("=")[1][:-1])
        speaker = parts[2]  # Полное имя спикера

        # Убираем строки, где время старта совпадает со временем стопа
        if start_time == stop_time:
            continue

        # Если список пуст или спикер изменился, добавляем новый интервал
        if not merged_segments or merged_segments[-1][2] != speaker:
            merged_segments.append([start_time, stop_time, speaker])
        else:
            # Если разница между stop предыдущего и start текущего <= 0.1, объединяем интервалы
            if abs(merged_segments[-1][1] - start_time) <= 0.1:
                merged_segments[-1][1] = stop_time
            else:
                merged_segments.append([start_time, stop_time, speaker])

    # Фильтруем интервалы, оставляя только те, у которых продолжительность > 1 сек
    filtered_segments = [seg for seg in merged_segments if seg[1] - seg[0] > 1.0]

    # Дополнительное объединение, если спикер совпадает (объединяем соседние интервалы)
    final_segments = []
    for seg in filtered_segments:
        if not final_segments or final_segments[-1][2] != seg[2]:
            final_segments.append(seg)
        else:
            final_segments[-1][1] = seg[1]  # Расширяем предыдущий интервал

    # Сохраняем результат
    with open(output_file, "w") as file:
        for start, stop, speaker in final_segments:
            file.write(f"start={start:.1f}s stop={stop:.1f}s {speaker}\n")
    
    print(f"Обработанные интервалы сохранены в {output_file}")

def convert_seconds_to_timestamp(seconds):
    """Преобразует секунды в формат HH:MM:SS,SSS"""
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    sec = int(seconds % 60)
    millisec = int((seconds % 1) * 1000)
    return f"{hours:02}:{minutes:02}:{sec:02},{millisec:03}"

def convert_diarization_file(input_path, output_path):
    """
    Читает файл диаризации, преобразует время в формат HH:MM:SS,SSS и сохраняет новый файл.
    """
    with open(input_path, "r") as file:
        lines = file.readlines()

    converted_lines = []
    for line in lines:
        match = re.match(r"start=([\d\.]+)s stop=([\d\.]+)s (speaker_SPEAKER_\d+)", line.strip())
        if match:
            start_seconds = float(match.group(1))
            stop_seconds = float(match.group(2))
            speaker = match.group(3)

            start_timestamp = convert_seconds_to_timestamp(start_seconds)
            stop_timestamp = convert_seconds_to_timestamp(stop_seconds)

            converted_lines.append(f"{start_timestamp} - {stop_timestamp} - {speaker}")

    with open(output_path, "w") as file:
        for line in converted_lines:
            file.write(line + "\n")

    print(f"Файл с преобразованным временем сохранён в {output_path}")
