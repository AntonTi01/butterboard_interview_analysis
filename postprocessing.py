import os
import re
from docx import Document

def clean_text(text: str) -> str:
    """Удаляет пустые строки между абзацами."""
    paragraphs = text.strip().split("\n")
    cleaned_paragraphs = [p.strip() for p in paragraphs if p.strip()]
    return "\n".join(cleaned_paragraphs)

def process_srt(input_path: str, output_path: str):
    """Обрабатывает .srt файл, удаляя пустые строки между субтитрами."""
    with open(input_path, "r", encoding="utf-8") as file:
        content = file.read()
    
    cleaned_content = clean_text(content)
    
    with open(output_path, "w", encoding="utf-8") as file:
        file.write(cleaned_content)
    print(f"Обработанный .srt файл сохранён в {output_path}")

def process_docx(input_path: str, output_path: str):
    """Обрабатывает .docx файл, удаляя пустые строки между абзацами."""
    doc = Document(input_path)
    cleaned_paragraphs = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            cleaned_paragraphs.append(text)
    
    new_doc = Document()
    for para in cleaned_paragraphs:
        new_doc.add_paragraph(para)
    
    new_doc.save(output_path)
    print(f"Обработанный .docx файл сохранён в {output_path}")

def process_transcription_file(input_path: str, output_path: str):
    """Обрабатывает отдельный файл .srt или .docx."""
    if input_path.endswith(".srt"):
        process_srt(input_path, output_path)
    elif input_path.endswith(".docx"):
        process_docx(input_path, output_path)
    else:
        print(f"Файл {input_path} не поддерживается (не .srt или .docx)")

if __name__ == "__main__":
    input_file = "results/raw/interview_1.docx"
    output_file = "results/postprocessed/interview_1_cleaned.docx"
    process_transcription_file(input_file, output_file)
